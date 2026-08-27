from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
md = (ROOT / 'PAPER.md').read_text(encoding='utf-8')
out = ROOT / 'PQ-ForensicVault-Research-Paper.docx'

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
styles['Normal'].font.size = Pt(10.5)
for name, size, color in [('Title', 24, '8F1D2C'), ('Heading 1', 17, '8F1D2C'), ('Heading 2', 13, 'B33A4A'), ('Heading 3', 11, '8F1D2C')]:
    styles[name].font.name = 'Aptos Display'
    styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos Display')
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text.strip())
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(rows):
    if not rows:
        return
    n = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], value, True, 'FFFFFF')
        shade(table.rows[0].cells[i], '8F1D2C')
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(n):
            set_cell_text(cells[i], row[i] if i < len(row) else '')
            if len(table.rows) % 2 == 0:
                shade(cells[i], 'F8EDEF')
    doc.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        if not line.strip() or set(line.strip()) <= set('|-: '):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

# Cover page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PQ-FORENSICVAULT')
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor(143,29,44)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Design and Evaluation of a Post-Quantum Chain-of-Custody Framework\nfor Classical Digital Evidence')
r.bold = True; r.font.size = Pt(19); r.font.color.rgb = RGBColor(55,55,55)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Reproducible Proof-of-Concept Study Using ECDSA-P256 and ML-DSA-65')
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(179,58,74)
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Research-paper draft\n').bold = True
p.add_run('Author: Goutham\n')
p.add_run('Benchmark environment: Node.js v22.13.0 · Linux x64\n')
p.add_run('Final repository tag: v2.2.2-readme-final\n')
p.add_run('27 August 2026')
doc.add_page_break()

# Content parser
lines = md.splitlines()
i = 0
while i < len(lines):
    line = lines[i]
    if not line.strip():
        i += 1; continue
    if line.startswith('---'):
        i += 1; continue
    if line.startswith('# '):
        # skip duplicate paper title after cover
        i += 1; continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=1); i += 1; continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=2); i += 1; continue
    if line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=3); i += 1; continue
    if line.startswith('|'):
        block=[]
        while i < len(lines) and lines[i].startswith('|'):
            block.append(lines[i]); i += 1
        add_table(parse_table(block)); continue
    if line.startswith('```'):
        code=[]; i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code.append(lines[i]); i += 1
        i += 1
        p=doc.add_paragraph(); p.style='No Spacing'
        r=p.add_run('\n'.join(code)); r.font.name='Consolas'; r.font.size=Pt(8.5)
        continue
    if line.startswith('!['):
        m=re.search(r'\(([^)]+)\)', line)
        if m and (ROOT / m.group(1)).exists():
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(ROOT / m.group(1)), width=Inches(6.2))
        i += 1; continue
    if line.startswith('*Figure'):
        p=doc.add_paragraph(line.strip('*')); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic=True; i += 1; continue
    if line.startswith('**') and line.endswith('**'):
        p=doc.add_paragraph(); r=p.add_run(line.strip('*')); r.bold=True; i += 1; continue
    if re.match(r'^\d+\. ', line) or line.startswith('- '):
        p=doc.add_paragraph(style='List Bullet' if line.startswith('- ') else 'List Number')
        p.add_run(re.sub(r'^\d+\. |^- ', '', line)); i += 1; continue
    p=doc.add_paragraph()
    # modest inline emphasis
    parts=re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold=True
        elif part.startswith('*') and part.endswith('*'):
            p.add_run(part[1:-1]).italic=True
        else:
            p.add_run(part)
    i += 1

# Header/footer
for sec in doc.sections:
    header = sec.header.paragraphs[0]
    header.text = 'PQ-ForensicVault  |  Research Paper'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(143,29,44)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('PQ-ForensicVault · Academic proof of concept · ')
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    footer._p.append(fld)
    for r in footer.runs: r.font.size = Pt(8)

doc.save(out)
print(out)
